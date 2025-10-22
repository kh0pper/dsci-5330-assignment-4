#!/usr/bin/env python3
"""
Create a professionally formatted DOCX memo based on README.md content
Following the style of Assignment2_Group3_Memo.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number_with_logo(section, logo_path):
    """Add Ford logo and page numbers to footer"""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Create a table in footer with 2 columns (logo left, page number right)
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False

    # Left cell for logo
    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(1.5)
    logo_para = left_cell.paragraphs[0]
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.0))

    # Right cell for page numbers
    right_cell = table.rows[0].cells[1]
    right_cell.width = Inches(5.0)
    page_para = right_cell.paragraphs[0]
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add page number field
    page_run = page_para.add_run()
    page_run.font.size = Pt(10)
    page_run.font.name = 'Calibri'
    page_run.font.color.rgb = RGBColor(0, 0, 91)  # Ford blue

    # Create page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)

    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_cover_page(doc, logo_path):
    """Create cover page with Ford branding"""
    # Add logo at top
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_path, width=Inches(2.5))

    # Add space
    doc.add_paragraph()
    doc.add_paragraph()

    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MANAGEMENT MODEL &\nSTRATEGY IMPLEMENTATION ANALYSIS")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 91)  # Ford blue
    title_run.font.name = 'Calibri'

    # Subtitle
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Ford Motor Company (2020-2025)")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0, 52, 120)
    subtitle_run.font.name = 'Calibri'

    # Add space
    for _ in range(5):
        doc.add_paragraph()

    # Add prepared by section
    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_run = prepared.add_run("Prepared by:\nBusiness Intelligence Team")
    prepared_run.font.size = Pt(14)
    prepared_run.font.name = 'Calibri'

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run("October 22, 2025")
    date_run.font.size = Pt(12)
    date_run.font.name = 'Calibri'

    # Add page break
    doc.add_page_break()

def parse_markdown_to_docx(doc, md_content):
    """Convert markdown content to DOCX with formatting"""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip metadata section at end
        if '**Memorandum Length:**' in line:
            break

        # Handle headings
        if line.startswith('# ') and not line.startswith('# MEMORANDUM'):
            # Skip the main title since we have cover page
            i += 1
            continue
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 91)
            heading.runs[0].font.name = 'Calibri'
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 52, 120)
            heading.runs[0].font.name = 'Calibri'
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=3)
            heading.runs[0].font.color.rgb = RGBColor(0, 80, 158)
            heading.runs[0].font.name = 'Calibri'

        # Handle memo header
        elif line.startswith('**TO:**') or line.startswith('**FROM:**') or line.startswith('**DATE:**') or line.startswith('**RE:**'):
            para = doc.add_paragraph()
            parts = line.split('**', 2)
            if len(parts) >= 3:
                label_run = para.add_run(parts[1])
                label_run.font.bold = True
                label_run.font.name = 'Calibri'
                label_run.font.size = Pt(11)
                value_run = para.add_run(parts[2])
                value_run.font.name = 'Calibri'
                value_run.font.size = Pt(11)

        # Handle tables
        elif line.startswith('|') and not in_table:
            in_table = True
            table_data = [line]

        elif line.startswith('|') and in_table:
            table_data.append(line)

        elif in_table and not line.startswith('|'):
            # End of table
            create_table_from_markdown(doc, table_data)
            in_table = False
            table_data = []
            i -= 1  # Re-process this line

        # Handle regular paragraphs
        elif line and not line.startswith('*[') and not line.startswith('[INSERT'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Process bold text
            process_formatted_text(para, line)

        i += 1

    # Handle any remaining table
    if in_table and table_data:
        create_table_from_markdown(doc, table_data)

def process_formatted_text(para, text):
    """Process text with bold and other formatting"""
    parts = []
    current = ""
    in_bold = False

    i = 0
    while i < len(text):
        if i < len(text) - 1 and text[i:i+2] == '**':
            if current:
                parts.append(('bold' if in_bold else 'normal', current))
                current = ""
            in_bold = not in_bold
            i += 2
        else:
            current += text[i]
            i += 1

    if current:
        parts.append(('bold' if in_bold else 'normal', current))

    for style, content in parts:
        run = para.add_run(content)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if style == 'bold':
            run.font.bold = True

def create_table_from_markdown(doc, table_lines):
    """Create a Word table from markdown table lines"""
    # Parse table
    rows = []
    for line in table_lines:
        if '---' in line:  # Skip separator line
            continue
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text

                # Format header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    # Set cell background color
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '00095B')  # Ford dark blue
                    cell._element.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()  # Add space after table

def main():
    """Main function to create the professional memo"""
    # Paths
    readme_path = '/home/kh0pp/dsci-5330-assignment-4/README.md'
    logo_path = '/home/kh0pp/dsci-5330-assignment-4/Ford_Motor_Company_Logo.png'
    output_path = '/home/kh0pp/dsci-5330-assignment-4/Ford_Management_Strategy_Memo.docx'

    # Read markdown content
    with open(readme_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Create cover page
    print("Creating cover page...")
    create_cover_page(doc, logo_path)

    # Add memo header
    print("Adding memo header...")
    memo_header = doc.add_heading('MEMORANDUM', level=1)
    memo_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    memo_header.runs[0].font.color.rgb = RGBColor(0, 0, 91)
    doc.add_paragraph()

    # Parse and add content
    print("Converting markdown content to DOCX...")
    parse_markdown_to_docx(doc, md_content)

    # Add footer with logo and page numbers to all sections except first (cover page)
    print("Adding footers with logo and page numbers...")
    for i, section in enumerate(doc.sections):
        if i > 0:  # Skip cover page
            add_page_number_with_logo(section, logo_path)

    # Save document
    print(f"Saving document to {output_path}...")
    doc.save(output_path)
    print("✓ Document created successfully!")

    return output_path

if __name__ == "__main__":
    output = main()
    print(f"\nProfessional memo saved to:\n{output}")
